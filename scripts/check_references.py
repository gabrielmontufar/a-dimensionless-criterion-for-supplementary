from __future__ import annotations

from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "manuscript" / "Demand_Polarity_Map_JEE_v40_final.docx"

EXPECTED_CITATION_TOKENS = [
    "Ancheta et al., 2014",
    "Allmond et al., 2015",
    "Ang and Tang, 2007",
    "ASCE/SEI, 2022",
    "Aviles and Perez-Rocha, 2003",
    "Bielak, 1975",
    "Boore, 2003",
    "Boore and Bommer, 2005",
    "CEN, 2004",
    "Chopra, 2017",
    "Clough and Penzien, 1993",
    "Dobry and Gazetas, 1986",
    "Elsabee and Morray, 1977",
    "FEMA, 2005",
    "FEMA, 2020",
    "Gazetas, 1983",
    "Gazetas, 1991",
    "Givens et al., 2016",
    "Jennings and Bielak, 1973",
    "Kausel, 2010",
    "Kausel and Roesset, 1975",
    "Kim and Stewart, 2003",
    "Kramer, 1996",
    "Luco and Westmann, 1971",
    "Melchers and Beck, 2018",
    "Montgomery, 2017",
    "Mylonakis and Gazetas, 2000",
    "Newmark and Rosenblueth, 1971",
    "NIST, 2012",
    "Novak and Beredugo, 1972",
    "Pais and Kausel, 1988",
    "Pitilakis et al., 2008",
    "Pitilakis et al., 2025a",
    "Pitilakis et al., 2025b",
    "Richart et al., 1970",
    "Roesset, 1980",
    "Saltelli et al., 2008",
    "Seed and Idriss, 1970",
    "Stewart et al., 1999",
    "Stewart et al., 2003",
    "Tao et al. (2024)",
    "Tileylioglu et al., 2011",
    "Veletsos and Meek, 1974",
    "Veletsos and Nair, 1975",
    "Wolf, 1985",
    "Wolf and Deeks, 2004",
    "Yang et al. (2024)",
    "Zhang and Far, 2024",
    "Gavras et al., 2023",
    "Hakhamaneshi et al., 2019",
    "NEES@UCSB",
]


def main() -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise SystemExit(f"python-docx is required for this check: {exc}") from exc

    doc = Document(str(DOCX))
    ref_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "References") + 1
    ref_end = next(
        (
            i
            for i in range(ref_start, len(doc.paragraphs))
            if doc.paragraphs[i].text.strip() == "Statements and Declarations"
        ),
        len(doc.paragraphs),
    )
    references = [p.text.strip() for p in doc.paragraphs[ref_start:ref_end] if p.text.strip()]
    body_parts = [p.text for p in doc.paragraphs[: ref_start - 1]]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                body_parts.append(cell.text)
    body = "\n".join(body_parts)
    reference_text = "\n".join(references)
    searchable_text = body + "\n" + reference_text
    def token_present(token: str) -> bool:
        if token in searchable_text:
            return True
        surname = token.split(" et al.")[0].split(",")[0].split(" and ")[0]
        return surname in searchable_text

    missing_tokens = [token for token in EXPECTED_CITATION_TOKENS if not token_present(token)]
    checks = {
        "reference_count_at_least_50": len(references) >= 50,
        "expected_citation_tokens_present": len(missing_tokens) == 0,
    }
    if not all(checks.values()):
        print(
            {
                "REFERENCE_CHECK": "FAIL",
                "checks": checks,
                "reference_count": len(references),
                "missing_tokens": missing_tokens,
            }
        )
        raise SystemExit(1)
    print({"REFERENCE_CHECK": "PASS", "reference_count": len(references)})


if __name__ == "__main__":
    main()
