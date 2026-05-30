from __future__ import annotations

from pathlib import Path
import re
from zipfile import ZipFile


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "manuscript" / "Demand_Polarity_Map_JEE_v40_final.docx"


def main() -> None:
    if not DOCX.exists():
        raise SystemExit(f"Missing DOCX: {DOCX}")
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise SystemExit(f"python-docx is required for this check: {exc}") from exc

    document = Document(str(DOCX))
    long_headings = []
    table_numbers = []
    forbidden_words = []
    def inspect_text(text: str, location: dict) -> None:
        if re.search(r"\b[Ii]dentifies\b", text):
            forbidden_words.append({**location, "word": "identifies"})

    for i, para in enumerate(document.paragraphs, start=1):
        style_name = para.style.name if para.style is not None else ""
        text = para.text.strip()
        if style_name.startswith("Heading") and len(text) > 160:
            long_headings.append({"paragraph": i, "style": style_name, "length": len(text)})
        match = re.match(r"^Table\s+(\d+)\.", text)
        if match:
            table_numbers.append(int(match.group(1)))
        inspect_text(text, {"paragraph": i})
    for ti, table in enumerate(document.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                inspect_text(cell.text.strip(), {"table": ti, "row": ri, "cell": ci})
    with ZipFile(DOCX) as zf:
        xml = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    checks = {
        "docx_exists": DOCX.exists(),
        "long_heading_count_zero": len(long_headings) == 0,
        "table_v2_absent": "Table V2" not in xml,
        "pending_marker_absent": "COMMIT_PENDING" not in xml,
        "legacy_version_labels_absent": all(label not in xml for label in ["v23", "v24", "v25"]),
        "table_numbers_unique": len(table_numbers) == len(set(table_numbers)),
        "table_numbers_monotonic": table_numbers == sorted(table_numbers),
        "identifies_absent": len(forbidden_words) == 0,
        "omml_equations_present": "<m:oMath" in xml,
    }
    if not all(checks.values()):
        print(
            {
                "DOCX_STYLE_CHECK": "FAIL",
                "checks": checks,
                "long_headings": long_headings,
                "table_numbers": table_numbers,
                "forbidden_words": forbidden_words,
            }
        )
        raise SystemExit(1)
    print({"DOCX_STYLE_CHECK": "PASS", "checks": checks})


if __name__ == "__main__":
    main()
